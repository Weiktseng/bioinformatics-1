#need instal pip bioython docx first
#python biopythonpaser.py CP003940.gb CP003940.fasta

import timeit, time                         #speed detect
stp=timeit.default_timer()


import concurrent.futures

#with concurrent.futures.ProcessPoolExecutor() as executor:

#----------------------------------------------------input gb----------------------------------------------------------------
#'''
from Bio import SeqIO
import sys

record = SeqIO.read(sys.argv[1], "genbank")
#print(record)
inner=record.features

#for c in range(3):
  #print(record.features[c])

name=[]
loc=[]
prod=[]
for f in record.features:                           #extract all date from gb and append to list
    if f.type != 'CDS': continue
    name.append(f.qualifiers['locus_tag'])           #tag:name,location:loc,product:prod
    loc.append(str(f.location))
    prod.append(f.qualifiers['product'])
    #print(f.type)
    #print(f.qualifiers['locus_tag'])
    #print(str(f.location))
    #print(f.qualifiers['product'])
    #print("----------------\n"); #, f)

#print('name:',name[0:4],'\n')
#print('location:',loc[0:4],'\n')
#print('product:',prod[0:4])
#print(type(loc[0][:-3]))
#print(type([1:2])
#print(len(name))
#print(loc)
#'''
#---------------------------------------------------input fasta--------------------------------------------------------------
#'''

genome=SeqIO.read(sys.argv[2], "fasta")
hseq=list(genome.seq)
     
import re
re.compile('\\d+')
re.compile('\\D+')

def LC(n,seq):                                               #/n every 50 bp
    seq1=''
    if  int(len(seq)/n) >0:  
        for c in range(int(len(seq)/n)):
            seq1=seq1+''.join(seq[c*n:c*n+n])+'\n'
        seq1=seq1+''.join(seq[(c+1)*n:])
    else:    
        seq1=''.join(seq)
    return seq1.strip('\n')

fi=open('seq_fasta.fasta','w')
no=0
for d in loc:
    ans = re.findall('\d+', d)
    seq2=hseq[int(ans[0]):int(ans[1])]
    #print(ans)
    #seq3=''.join(seq2)
    
    fi.write(str(name[no][0]))   
    no=no+1 
    fi.write('\n') 
    fi.write(LC(50,seq2))
    fi.write('\n')
    fi.write('.'*50+'\n')
fi.close()


#'''
#----------------------------------------------------out put to excel-----------------------------------------------
#'''   

#Program:
#This program will record log into Excel.
#History:
#20170707 Kuanlin Chen


#匯入模組(Module)

import xlwt

#建立Workbook物件
book = xlwt.Workbook(encoding="utf-8")
#使用Workbook裡的add_sheet函式來建立Worksheet
sheet1 = book.add_sheet("Sheet1")

def main(orig_args):
    filename = "seq_fasta.xls"
    output(filename)

def output(filename):
    #使用Worksheet裡的write函式將值寫入
    sheet1.write(0,0,'Name')
    sheet1.write(0,1,'location')
    sheet1.write(0,2,'product')

    stime = time.time()
    for c in range(len(name)):
        sheet1.write(c+1,0,name[c])
        sheet1.write(c+1,1,loc[c])
        sheet1.write(c+1,2,prod[c])
        #if c % 250 == 0: print("Time to write to Excel: %s" % round(time.time() - stime, 2))
    print("Time to write to Excel: %s" % round(time.time() - stime, 2))

    #將Workbook儲存為原生Excel格式的檔案
    book.save(filename)

if __name__ == '__main__':
    main(sys.argv)  

#'''
#----------------------------------------------------out put to word-----------------------------------------------
#'''
import urllib.request
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('CDS list', 0)

p = document.add_paragraph('')
p.add_run('group9').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True

document.add_heading('results', level=1)
document.add_paragraph('')

#document.add_paragraph('first item in unordered list')
#document.add_paragraph('first item in ordered list')

#urllib.request.urlretrieve("http://placehold.it/350x150", "placeholder.png")
#document.add_picture('placeholder.png', width=Inches(1.25))

# To speed up writing the table, see: https://github.com/python-openxml/python-docx/issues/174
# COLUMNS = 7
# table = document.add_table(rows=1000, columns=COLUMNS)
# table_cells = table._cells
# for i in range(ROWS):
#     row_cells = table_cells[i*COLUMNS:(i+1)*COLUMNS]
#     #Add text to row_cells
# I used the technique above and the time to create and populate 
# the table dropped from about 70 seconds to 2 seconds! Thank you stumpyyy!.

table = document.add_table(rows=len(name), cols=3,style="Medium Shading 1 Accent 1")
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'name'
hdr_cells[1].text = 'location'
hdr_cells[2].text = 'product'

stime = time.time()

table_cells = table._cells
for c in range(len(name)):
    #row_cells=table._cells[c*3:(c+1)*3]  #do this in loop will cost 585.1576336885572 sec
    row_cells=table_cells[c*3:(c+1)*3]           
    row_cells[0].text=name[c] # table_cells[c*3  ].text = name[c]
    row_cells[1].text=loc[c]  # table_cells[c*3+1].text = loc[c]
    row_cells[2].text=prod[c] # table_cells[c*3+2].text = prod[c]
    # table.cell(c+1, 0).text = name[c]
    # table.cell(c+1, 1).text = loc[c]
    # table.cell(c+1, 2).text = prod[c]
    if c % 250 == 0: print("Time to write to Word: %s" % round(time.time() - stime, 2))
print("Time to write to Word: %s" % round(time.time() - stime, 2))

print(table._cells[0:3])
document.add_page_break()
document.save('CDS_list.docx')

endp=timeit.default_timer()
print('holoprogram time cost',endp-stp,'sec')
#'''
#-----------------------------------------------------finish-------------------------------------------------------

'''
#SeqIO.read structure:
  #.annotations
  #.dbxrefs
  #.description
  #.id
  #.seq
  #.features[n]:        it's a big list that each element contain type,location and qualifiers 
    #.type              include CDS,GENE
    #.location
    #.qualifiers['Key']:       it's a ordereddict. use Key to search

    Key: codon_start, Value: ['1']
    Key: db_xref, Value: ['GI:45478716', 'GeneID:2767712']
    Key: gene, Value: ['pim']
    Key: locus_tag, Value: ['YP_pPCP05']
    Key: note, Value: ['similar to many previously sequenced pesticin immunity ...']
    Key: product, Value: ['pesticin immunity protein']
    Key: protein_id, Value: ['NP_995571.1']
    Key: transl_table, Value: ['11']
    Key: translation, Value: ['MGGGMISKLFCLALIFLSSSGLAEKNTYTAKDILQNLELNTFGNSLSH...']
'''
    